# Étape 1 : Installer les dépendances
#pip install faker reportlab python-docx pandas googletrans==3.1.0a0

# Étape 2 : Importer les bibliothèques
import random
import os
import zipfile
import uuid
import pandas as pd
from faker import Faker
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from googletrans import Translator
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime

# Étape 3 : Initialiser Faker pour différentes langues
fake_fr = Faker('fr_FR')
fake_en = Faker('en_US')
fake_ru = Faker('ru_RU')
fake_it = Faker('it_IT')

# Initialiser le traducteur
translator = Translator()

# Listes de prénoms et noms djiboutiens
djiboutian_first_names = [
    # Masculins
    "Ahmed", "Hassan", "Mahad", "Ismail", "Abdourahman", "Ali", "Ibrahim", "Abdi", "Youssouf", "Souleiman",
    "Jamal", "Zakaria", "Mohamed", "Farhan", "Idriss", "Faysal", "Yassin", "Bilal", "Omar", "Abdallah",
    "Barkhad", "Anwar", "Saïd", "Nour", "Dini", "Mahamoud", "Khaled", "Awaleh", "Bileh", "Muse",
    "Saad", "Salah", "Warsame", "Guedi", "Walid", "Abdirahman", "Haroun", "Moustapha", "Abshir",
    "Suleyman", "Dawoud", "Houssein", "Abokor", "Aweis", "Hersi", "Aaden", "Ayub", "Farah", "Ayman",
    "Rachid", "Madar", "Nuradin", "Hassanali", "Kalif", "Deria", "Saidou", "Samatar", "Tahlil",
    "Wais", "Darar", "Robleh", "Barre", "Aliyow", "Djamal", "Sharmarke", "Kamil", "Mahdi", "Gedi",
    "Duale", "Barkat",
    # Féminins
    "Fadumo", "Ayan", "Hibo", "Ifrah", "Nimo", "Hawa", "Sagal", "Asha", "Safia", "Hodan",
    "Zamzam", "Amal", "Faiza", "Rahma", "Amina", "Khadija", "Naima", "Yasmin", "Salma", "Ikram",
    "Roda", "Halima", "Saida", "Habiba", "Mariam", "Fatouma", "Asli", "Souad", "Batula", "Anisa",
    "Hani", "Deeqa", "Shaafi", "Sahra", "Rihana", "Munira", "Bilan", "Layla", "Nadra", "Muna",
    "Mako", "Sucaad", "Faisa", "Anab", "Haniyo", "Sulekha", "Fawzia", "Farhia", "Ifra", "Sudi",
    "Lubna", "Bashira", "Nura", "Zamzama", "Zaynab", "Deka", "Ikran", "Saynab", "Jawahir",
    "Farhiyo", "Ubah", "Shukri", "Ismahan", "Barwaqo", "Rukia", "Misra", "Nasteexo", "Samia",
    "Asma", "Rahimo"
]
djiboutian_last_names = [
    "Abdi Farah", "Ali Robleh", "Hassan Dini", "Youssouf Omar", "Mohamed Farhan", "Ismail Awaleh",
    "Ahmed Gedi", "Ibrahim Duale", "Mahad Saïd", "Abdirahman Muse", "Farhan Mahamoud", "Guedi Warsame",
    "Omar Abdi", "Abdourahman Walid", "Khaled Awaleh", "Abshir Robleh", "Suleyman Barre",
    "Djamal Samatar", "Robleh Tahlil", "Barkhad Robleh", "Souleiman Abdi", "Saïd Deria",
    "Warsame Madar", "Nour Hassan", "Mahdi Robleh", "Dawoud Robleh", "Abokor Wais", "Deria Saïd",
    "Saad Abdi", "Houssein Hassan", "Moustapha Abdi", "Walid Barkhad", "Aliyow Djamal",
    "Abdi Sharmarke", "Salah Djamal", "Ayub Farah", "Rachid Madar", "Mahamoud Saidou",
    "Awaleh Barkat", "Yassin Madar", "Abdi Barre", "Barkat Kalif", "Madar Guedi", "Jamal Sharmarke",
    "Bilal Saïd", "Robleh Wais", "Barre Awaleh", "Khalif Warsame", "Farah Hersi", "Abdallah Madar",
    "Hersi Samatar", "Nouradin Tahlil", "Saidou Hassan", "Kalif Dawoud", "Walid Saidou",
    "Djamal Barkat", "Barkat Awaleh", "Faysal Barre", "Samatar Guedi", "Deria Hersi",
    "Aweis Tahlil", "Robleh Barkhad", "Awaleh Saidou", "Ali Tahlil", "Abdirahman Barkhad",
    "Madar Warsame", "Ismail Deria", "Omar Abdi", "Jamal Saidou", "Saïd Abshir", "Ali Robleh",
    "Youssouf Madar", "Barkhad Muse", "Tahlil Warsame", "Barkat Duale", "Awaleh Guedi",
    "Samatar Mahdi", "Walid Nour", "Muse Saïd", "Saidou Barkat", "Mahamoud Djamal",
    "Sharmarke Warsame", "Deria Robleh", "Guedi Awaleh", "Hassan Kalif"
]

# Noms étrangers spécifiques par nationalité
foreign_first_names = {
    "French": ["Lucas", "Hugo", "Enzo", "Arthur", "Léo", "Nathan", "Raphaël", "Jules", "Tom", "Clément",
               "Louis", "Adam", "Noé", "Théo", "Maxime"],
    "English": ["James", "John", "Robert", "Michael", "William", "David", "Richard", "Charles", "Joseph",
                "Thomas", "Daniel", "Matthew", "Andrew", "Kevin", "Brian"],
    "Russian": ["Ivan", "Dimitri", "Alexei", "Nikolai", "Mikhail", "Vladimir", "Yuri", "Andrei", "Sergei",
                "Pavel"],
    "Italian": ["Luca", "Matteo", "Marco", "Andrea", "Alessandro", "Gabriele", "Giovanni", "Antonio",
                "Stefano", "Francesco"],
    "Ethiopian": ["Bekele", "Tadesse", "Tesfaye", "Abebe", "Dawit", "Getachew", "Hailu", "Kassahun",
                  "Solomon", "Yosef"],
    "Other": ["Ali", "Omar", "Amir", "Malik", "Ibrahim", "Hamza", "Sami", "Elias", "Youssef", "Karim"]
}
foreign_last_names = [
    "Dubois", "Lefèvre", "Bernard", "Girard", "Morel",
    "Smith", "Johnson", "Williams", "Brown", "Jones",
    "Petrov", "Sokolov", "Kuznetsov", "Ivanov", "Makarov",
    "Rossi", "Bianchi", "Romano", "Conti", "Gallo",
    "Teshome", "Gebre", "Mekonnen", "Alemu", "Desta",
    "El Fassi", "Ben Omar", "Diouf", "Bamba", "Haddad"
]

# Liste des 200 métiers
jobs = [
    "Enseignant", "Médecin", "Ingénieur informatique", "Infirmier", "Commerçant", "Comptable", "Avocat",
    "Chauffeur", "Électricien", "Pharmacien", "Journaliste", "Développeur web", "Cuisinier", "Secrétaire",
    "Logisticien", "Pêcheur", "Agriculteur", "Architecte", "Dentiste", "Vétérinaire", "Banquier", "Marketeur",
    "Plombier", "Mécanicien", "Coiffeur", "Tailleur", "Boulanger", "Pâtissier", "Serveur", "Barman",
    "Agent immobilier", "Agent de sécurité", "Policier", "Pompier", "Militaire", "Technicien de maintenance",
    "Analyste financier", "Consultant", "Graphiste", "Photographe", "Vidéaste", "Monteur vidéo", "Rédacteur",
    "Traducteur", "Interprète", "Bibliothécaire", "Archiviste", "Conservateur de musée", "Guide touristique",
    "Agent de voyage", "Pilote", "Hôtesse de l'air", "Contrôleur aérien", "Marin", "Capitaine de navire",
    "Paysagiste", "Jardinier", "Fleuriste", "Écologiste", "Biologiste", "Chimiste", "Physicien", "Mathématicien",
    "Statisticien", "Data scientist", "Data analyst", "Ingénieur civil", "Ingénieur mécanique", "Ingénieur électrique",
    "Urbaniste", "Géologue", "Météorologue", "Océanographe", "Astronome", "Psychologue", "Psychiatre",
    "Travailleur social", "Conseiller d'orientation", "Coach", "Formateur", "Animateur", "Éducateur spécialisé",
    "Orthophoniste", "Kinésithérapeute", "Ergothérapeute", "Opticien", "Audioprothésiste", "Sage-femme",
    "Aide-soignant", "Ambulancier", "Laborantin", "Chercheur", "Professeur d'université", "Assistant de recherche",
    "Économiste", "Sociologue", "Anthropologue", "Historien", "Géographe", "Politicien", "Diplomate",
    "Agent des douanes", "Inspecteur fiscal", "Auditeur", "Actuaire", "Assureur", "Courtier", "Notaire",
    "Huissier", "Juge", "Procureur", "Greffier", "Planificateur", "Gestionnaire de projet", "Scrum master",
    "Product owner", "UX designer", "UI designer", "Développeur mobile", "Développeur backend", "Développeur frontend",
    "Administrateur système", "Administrateur réseau", "Spécialiste cybersécurité", "Testeur QA", "Technicien support",
    "Community manager", "Responsable communication", "Publicitaire", "Organisateur d'événements", "Organisateur de mariage",
    "DJ", "Musicien", "Chanteur", "Acteur", "Réalisateur", "Scénariste", "Producteur", "Costumier", "Maquilleur",
    "Décorateur", "Éclairagiste", "Technicien son", "Technicien lumière", "Styliste", "Modéliste", "Couturier",
    "Bijoutier", "Horloger", "Ébéniste", "Menuisier", "Charpentier", "Maçon", "Peintre en bâtiment", "Carreleur",
    "Couvreur", "Vitrier", "Serrurier", "Forgeron", "Soudeur", "Métallurgiste", "Conducteur d'engins", "Grutier",
    "Chauffeur poids lourd", "Livreur", "Magasinier", "Employé de banque", "Caissier", "Vendeur", "Manager de magasin",
    "Merchandiser", "Sommelier", "Œnologue", "Viticulteur", "Brasseur", "Distillateur", "Apiculteur", "Éleveur",
    "Vacher", "Berger", "Maraîcher", "Arboriculteur", "Pépiniériste", "Semencier", "Paysan", "Technicien agricole",
    "Vulcanologue", "Sismologue", "Hydrologue", "Glaciologue", "Cartographe", "Topographe", "Géomaticien",
    "Épidémiologiste", "Généticien", "Biochimiste", "Microbiologiste", "Pharmacologue", "Toxicologue",
    "Nutritionniste", "Diététicien", "Préparateur physique", "Entraîneur", "Arbitre", "Agent sportif",
    "Journaliste sportif", "Commentateur", "Éditeur", "Libraire", "Imprimeur", "Relieur"
]

# Compétences spécifiques par métier (5 à 15 par métier)
job_skills = {
    "Enseignant": ["Pédagogie", "Gestion de classe", "Planification de cours", "Évaluation des élèves", "Technologies éducatives", "Communication", "Adaptabilité", "Créativité", "Travail d'équipe", "Motivation des élèves"],
    "Médecin": ["Diagnostic médical", "Soins aux patients", "Pharmacologie", "Chirurgie", "Gestion des urgences", "Empathie", "Analyse de dossiers médicaux", "Communication patient", "Éthique médicale", "Mise à jour des connaissances"],
    "Ingénieur informatique": ["Programmation", "Résolution de problèmes", "Conception de systèmes", "Gestion de bases de données", "Cybersécurité", "Cloud computing", "DevOps", "Analyse des besoins", "Test de logiciels", "Maintenance système", "Python", "Java", "Agilité"],
    "Infirmier": ["Soins infirmiers", "Administration de médicaments", "Surveillance des patients", "Premiers secours", "Gestion des dossiers médicaux", "Empathie", "Communication", "Travail sous pression", "Hygiène hospitalière", "Collaboration médicale"],
    "Commerçant": ["Négociation", "Gestion des stocks", "Service client", "Vente", "Marketing", "Comptabilité de base", "Analyse de marché", "Fidélisation client", "Gestion du temps", "Résolution de conflits"],
    "Comptable": ["Comptabilité générale", "Analyse financière", "Fiscalité", "Logiciels comptables", "Budgétisation", "Audit", "Reporting", "Gestion de la paie", "Conformité réglementaire", "Excel avancé"],
    "Avocat": ["Droit", "Plaidoyer", "Rédaction juridique", "Négociation", "Recherche juridique", "Éthique professionnelle", "Gestion de dossiers", "Communication client", "Résolution de litiges", "Connaissance des lois"],
    "Chauffeur": ["Conduite sécuritaire", "Navigation GPS", "Entretien du véhicule", "Respect des horaires", "Service client", "Connaissance des routes", "Gestion du stress", "Premiers secours", "Règles de circulation"],
    "Électricien": ["Installation électrique", "Dépannage", "Lecture de schémas", "Normes de sécurité", "Maintenance", "Câblage", "Test de circuits", "Utilisation d'outils électriques", "Connaissance des normes"],
    "Pharmacien": ["Pharmacologie", "Dispensation de médicaments", "Conseil aux patients", "Gestion des stocks", "Réglementation pharmaceutique", "Communication", "Analyse des prescriptions", "Vigilance médicamenteuse"],
    "Journaliste": ["Rédaction", "Recherche d'informations", "Interview", "Éthique journalistique", "Montage vidéo", "Réseaux sociaux", "Analyse de données", "Communication", "Photographie"],
    "Développeur web": ["HTML/CSS", "JavaScript", "Frameworks (React, Angular)", "Backend (Node.js, Django)", "Bases de données", "SEO", "Tests unitaires", "Git", "UX/UI", "Responsive design"],
    "Cuisinier": ["Cuisine", "Gestion des ingrédients", "Hygiène alimentaire", "Créativité culinaire", "Gestion du temps", "Service client", "Menu planning", "Techniques de cuisson"],
    "Secrétaire": ["Gestion d'agenda", "Rédaction de courriers", "Accueil", "Organisation", "Informatique (Office)", "Communication", "Classement", "Prise de notes"],
    "Logisticien": ["Gestion de la chaîne d'approvisionnement", "Planification logistique", "Gestion des stocks", "Transport", "Négociation avec fournisseurs", "Analyse des coûts", "ERP", "Coordination"],
    "Pêcheur": ["Pêche artisanale", "Navigation maritime", "Entretien des filets", "Connaissance des espèces", "Sécurité en mer", "Réparation de bateaux", "Gestion des captures", "Météorologie"],
    "Agriculteur": ["Culture agricole", "Gestion des sols", "Irrigation", "Récolte", "Utilisation de machines agricoles", "Connaissance des semences", "Protection des cultures", "Comptabilité agricole"],
    "Architecte": ["Conception architecturale", "CAO (AutoCAD)", "Gestion de projets", "Normes de construction", "Urbanisme", "Rendu 3D", "Communication client", "Budgétisation", "Durabilité"],
    "Dentiste": ["Soins dentaires", "Chirurgie dentaire", "Diagnostic oral", "Radiologie dentaire", "Prothèses", "Hygiène buccale", "Communication patient", "Gestion de cabinet"],
    "Vétérinaire": ["Soins animaux", "Chirurgie vétérinaire", "Diagnostic", "Pharmacologie vétérinaire", "Communication client", "Gestion de clinique", "Bien-être animal"],
    "Banquier": ["Gestion de comptes", "Analyse de crédit", "Conseil financier", "Conformité réglementaire", "Vente de produits bancaires", "Communication client", "Analyse de risques"],
    "Marketeur": ["Stratégie marketing", "Analyse de marché", "Publicité", "Réseaux sociaux", "SEO/SEM", "Gestion de campagnes", "Communication", "Créativité"],
    "Plombier": ["Installation sanitaire", "Réparation de fuites", "Lecture de plans", "Soudure", "Maintenance", "Normes de sécurité", "Service client"],
    "Mécanicien": ["Réparation de véhicules", "Diagnostic mécanique", "Entretien", "Utilisation d'outils", "Connaissance des moteurs", "Normes de sécurité", "Service client"],
    "Coiffeur": ["Coupe de cheveux", "Coloration", "Coiffure", "Conseil client", "Hygiène", "Gestion du salon", "Créativité"],
    "Tailleur": ["Couture", "Prise de mesures", "Conception de vêtements", "Réparation", "Conseil client", "Gestion des tissus", "Précision"],
    "Boulanger": ["Pétrissage", "Cuisson", "Gestion des ingrédients", "Hygiène alimentaire", "Créativité", "Service client", "Gestion du temps"],
    "Pâtissier": ["Pâtisserie", "Décoration", "Gestion des ingrédients", "Hygiène alimentaire", "Créativité", "Précision", "Service client"],
    "Serveur": ["Service client", "Prise de commandes", "Gestion du temps", "Communication", "Hygiène", "Travail d'équipe", "Résolution de problèmes"],
    "Barman": ["Mixologie", "Service client", "Gestion des stocks", "Hygiène", "Communication", "Créativité", "Travail sous pression"],
    "Agent immobilier": ["Négociation", "Évaluation immobilière", "Vente", "Conseil client", "Connaissance du marché", "Rédaction de contrats", "Communication"],
    "Agent de sécurité": ["Surveillance", "Gestion des conflits", "Premiers secours", "Connaissance des lois", "Communication", "Vigilance", "Rapport"],
    "Policier": ["Application de la loi", "Investigation", "Sécurité publique", "Communication", "Premiers secours", "Gestion des conflits", "Éthique"],
    "Pompier": ["Extinction d'incendies", "Sauvetage", "Premiers secours", "Maintenance d'équipements", "Travail d'équipe", "Gestion du stress"],
    "Militaire": ["Stratégie", "Combat", "Discipline", "Premiers secours", "Communication", "Physique", "Maintenance d'équipements"],
    "Technicien de maintenance": ["Dépannage", "Maintenance préventive", "Lecture de schémas", "Utilisation d'outils", "Normes de sécurité", "Rapport"],
    "Analyste financier": ["Analyse de données", "Modélisation financière", "Évaluation des risques", "Excel avancé", "Rapport", "Connaissance des marchés"],
    "Consultant": ["Analyse de problèmes", "Conseil stratégique", "Communication client", "Gestion de projets", "Recherche", "Présentation"],
    "Graphiste": ["Conception graphique", "Photoshop", "Illustrator", "Créativité", "Communication client", "Gestion de projets"],
    "Photographe": ["Photographie", "Retouche d'images", "Éclairage", "Créativité", "Communication client", "Gestion d'équipements"],
    "Vidéaste": ["Tournage", "Montage vidéo", "Éclairage", "Son", "Créativité", "Logiciels de montage"],
    "Monteur vidéo": ["Montage", "Effets spéciaux", "Synchronisation audio", "Logiciels (Premiere, Final Cut)", "Créativité"],
    "Rédacteur": ["Rédaction", "Recherche", "SEO", "Correction", "Communication", "Créativité"],
    "Traducteur": ["Traduction", "Connaissance linguistique", "Recherche terminologique", "Précision", "Gestion du temps"],
    "Interprète": ["Interprétation simultanée", "Connaissance linguistique", "Communication", "Écoute active", "Gestion du stress"],
    "Bibliothécaire": ["Gestion de collections", "Recherche documentaire", "Service client", "Informatique", "Organisation"],
    "Archiviste": ["Gestion d'archives", "Numérisation", "Recherche", "Organisation", "Connaissance historique"],
    "Conservateur de musée": ["Gestion de collections", "Conservation", "Recherche", "Organisation d'expositions", "Communication"],
    "Guide touristique": ["Connaissance culturelle", "Communication", "Organisation", "Multilinguisme", "Service client"],
    "Agent de voyage": ["Planification de voyages", "Réservations", "Conseil client", "Connaissance des destinations", "Communication"],
    "Pilote": ["Pilotage", "Navigation", "Sécurité aérienne", "Communication", "Gestion du stress", "Maintenance"],
    "Hôtesse de l'air": ["Service client", "Sécurité aérienne", "Premiers secours", "Communication", "Multilinguisme"],
    "Contrôleur aérien": ["Gestion du trafic aérien", "Communication", "Prise de décision", "Gestion du stress", "Connaissance des protocoles"],
    "Marin": ["Navigation", "Maintenance de navires", "Sécurité maritime", "Communication", "Travail d'équipe"],
    "Capitaine de navire": ["Commandement", "Navigation", "Sécurité maritime", "Gestion d'équipe", "Maintenance"],
    "Paysagiste": ["Conception de jardins", "Entretien paysager", "Connaissance des plantes", "CAO", "Communication client"],
    "Jardinier": ["Entretien des plantes", "Taille", "Irrigation", "Connaissance des sols", "Utilisation d'outils"],
    "Fleuriste": ["Composition florale", "Connaissance des fleurs", "Service client", "Créativité", "Gestion des stocks"],
    "Écologiste": ["Recherche environnementale", "Analyse de données", "Conservation", "Communication", "Sensibilisation"],
    "Biologiste": ["Recherche scientifique", "Analyse de données", "Expérimentation", "Rédaction scientifique", "Connaissance biologique"],
    "Chimiste": ["Analyse chimique", "Synthèse", "Expérimentation", "Sécurité en laboratoire", "Rédaction scientifique"],
    "Physicien": ["Recherche scientifique", "Modélisation", "Expérimentation", "Analyse de données", "Mathématiques"],
    "Mathématicien": ["Modélisation", "Analyse", "Résolution de problèmes", "Programmation", "Enseignement"],
    "Statisticien": ["Analyse de données", "Modélisation statistique", "Logiciels (R, SPSS)", "Rapport", "Interprétation"],
    "Data scientist": ["Machine learning", "Analyse de données", "Programmation (Python, R)", "Visualisation", "Statistiques"],
    "Data analyst": ["Analyse de données", "Visualisation", "SQL", "Excel", "Rapport"],
    "Ingénieur civil": ["Conception de structures", "Gestion de projets", "CAO", "Normes de construction", "Budgétisation"],
    "Ingénieur mécanique": ["Conception mécanique", "CAO", "Tests", "Maintenance", "Gestion de projets"],
    "Ingénieur électrique": ["Conception de circuits", "Maintenance", "Normes de sécurité", "CAO", "Gestion de projets"],
    "Urbaniste": ["Planification urbaine", "Analyse spatiale", "Consultation publique", "SIG", "Réglementation"],
    "Géologue": ["Analyse des sols", "Cartographie", "Recherche", "Évaluation des risques", "Rapport"],
    "Météorologue": ["Prévision météo", "Analyse de données", "Modélisation", "Communication", "Instruments météo"],
    "Océanographe": ["Recherche marine", "Analyse de données", "Échantillonnage", "Modélisation", "Conservation"],
    "Astronome": ["Observation", "Analyse de données", "Modélisation", "Recherche", "Programmation"],
    "Psychologue": ["Écoute active", "Diagnostic", "Conseil", "Recherche", "Communication"],
    "Psychiatre": ["Diagnostic", "Prescription", "Thérapie", "Communication patient", "Éthique"],
    "Travailleur social": ["Accompagnement", "Évaluation des besoins", "Médiation", "Communication", "Empathie"],
    "Conseiller d'orientation": ["Évaluation", "Conseil", "Connaissance du marché", "Communication", "Planification"],
    "Coach": ["Motivation", "Planification", "Écoute", "Fixation d'objectifs", "Communication"],
    "Formateur": ["Pédagogie", "Conception de cours", "Animation", "Évaluation", "Communication"],
    "Animateur": ["Animation", "Organisation", "Communication", "Créativité", "Travail d'équipe"],
    "Éducateur spécialisé": ["Accompagnement", "Pédagogie", "Évaluation", "Communication", "Empathie"],
    "Orthophoniste": ["Diagnostic", "Thérapie", "Communication", "Évaluation", "Pédagogie"],
    "Kinésithérapeute": ["Rééducation", "Massage", "Évaluation", "Communication patient", "Anatomie"],
    "Ergothérapeute": ["Réadaptation", "Évaluation", "Planification", "Communication", "Empathie"],
    "Opticien": ["Évaluation visuelle", "Fabrication de lunettes", "Conseil client", "Précision", "Service client"],
    "Audioprothésiste": ["Évaluation auditive", "Ajustement d'appareils", "Conseil client", "Communication", "Technologie"],
    "Sage-femme": ["Accouchement", "Suivi prénatal", "Conseil", "Communication", "Premiers secours"],
    "Aide-soignant": ["Soins de base", "Hygiène", "Communication patient", "Empathie", "Travail d'équipe"],
    "Ambulancier": ["Transport médical", "Premiers secours", "Conduite sécuritaire", "Communication", "Gestion du stress"],
    "Laborantin": ["Analyse", "Prélèvements", "Sécurité en laboratoire", "Précision", "Rapport"],
    "Chercheur": ["Recherche", "Analyse de données", "Rédaction scientifique", "Expérimentation", "Collaboration"],
    "Professeur d'université": ["Enseignement", "Recherche", "Publication", "Mentorat", "Conception de cours"],
    "Assistant de recherche": ["Collecte de données", "Analyse", "Rédaction", "Collaboration", "Organisation"],
    "Économiste": ["Analyse économique", "Modélisation", "Prévision", "Rapport", "Connaissance des marchés"],
    "Sociologue": ["Recherche", "Analyse qualitative", "Analyse quantitative", "Rédaction", "Observation"],
    "Anthropologue": ["Recherche culturelle", "Observation", "Analyse", "Rédaction", "Terrain"],
    "Historien": ["Recherche historique", "Analyse de sources", "Rédaction", "Enseignement", "Archivage"],
    "Géographe": ["Analyse spatiale", "Cartographie", "SIG", "Recherche", "Rédaction"],
    "Politicien": ["Prise de décision", "Communication", "Négociation", "Stratégie", "Connaissance des lois"],
    "Diplomate": ["Négociation", "Communication", "Connaissance internationale", "Multilinguisme", "Protocole"],
    "Agent des douanes": ["Contrôle", "Réglementation", "Communication", "Vigilance", "Rapport"],
    "Inspecteur fiscal": ["Audit fiscal", "Analyse", "Connaissance des lois", "Rapport", "Communication"],
    "Auditeur": ["Audit", "Analyse financière", "Rapport", "Conformité", "Communication"],
    "Actuaire": ["Analyse des risques", "Modélisation", "Statistiques", "Rapport", "Logiciels actuariels"],
    "Assureur": ["Évaluation des risques", "Vente", "Conseil client", "Rédaction de polices", "Communication"],
    "Courtier": ["Négociation", "Conseil", "Connaissance des marchés", "Communication client", "Analyse"],
    "Notaire": ["Rédaction d'actes", "Conseil juridique", "Connaissance des lois", "Communication", "Éthique"],
    "Huissier": ["Exécution des décisions", "Rédaction", "Connaissance juridique", "Communication", "Organisation"],
    "Juge": ["Interprétation des lois", "Prise de décision", "Éthique", "Communication", "Analyse"],
    "Procureur": ["Poursuite", "Analyse juridique", "Plaidoyer", "Éthique", "Communication"],
    "Greffier": ["Gestion des dossiers", "Rédaction", "Organisation", "Connaissance juridique", "Communication"],
    "Planificateur": ["Planification", "Analyse", "Coordination", "Communication", "Gestion de projets"],
    "Gestionnaire de projet": ["Planification", "Coordination", "Budgétisation", "Communication", "Résolution de problèmes"],
    "Scrum master": ["Facilitation", "Agilité", "Coordination", "Communication", "Résolution de conflits"],
    "Product owner": ["Définition de produit", "Priorisation", "Communication", "Agilité", "Analyse des besoins"],
    "UX designer": ["Recherche utilisateur", "Prototypage", "Tests d'utilisabilité", "Design d'interface", "Communication"],
    "UI designer": ["Design d'interface", "Prototypage", "Graphisme", "Outils (Figma)", "Créativité"],
    "Développeur mobile": ["Programmation (Swift, Kotlin)", "Tests", "Conception d'applications", "Git", "API"],
    "Développeur backend": ["Programmation (Python, Java)", "Bases de données", "API", "Sécurité", "Git"],
    "Développeur frontend": ["HTML/CSS", "JavaScript", "Frameworks (React)", "Responsive design", "Git"],
    "Administrateur système": ["Gestion de serveurs", "Sécurité", "Réseaux", "Scripting", "Maintenance"],
    "Administrateur réseau": ["Configuration réseau", "Sécurité", "Dépannage", "Monitoring", "Protocoles"],
    "Spécialiste cybersécurité": ["Analyse des menaces", "Pentesting", "Sécurité réseau", "Cryptographie", "Rapport"],
    "Testeur QA": ["Tests fonctionnels", "Tests automatisés", "Rédaction de cas de test", "Bug tracking", "Communication"],
    "Technicien support": ["Dépannage", "Support utilisateur", "Connaissance IT", "Communication", "Documentation"],
    "Community manager": ["Gestion des réseaux sociaux", "Création de contenu", "Engagement", "Analyse", "Communication"],
    "Responsable communication": ["Stratégie de communication", "Rédaction", "Gestion de crise", "Relations publiques", "Coordination"],
    "Publicitaire": ["Conception de campagnes", "Analyse de marché", "Créativité", "Communication", "Gestion de projets"],
    "Organisateur d'événements": ["Planification", "Coordination", "Budgétisation", "Communication", "Négociation"],
    "Organisateur de mariage": ["Planification", "Coordination", "Créativité", "Communication client", "Budgétisation"],
    "DJ": ["Mixage", "Connaissance musicale", "Animation", "Équipements audio", "Créativité"],
    "Musicien": ["Interprétation", "Composition", "Connaissance musicale", "Pratique d'instruments", "Créativité"],
    "Chanteur": ["Chant", "Interprétation", "Scène", "Connaissance musicale", "Créativité"],
    "Acteur": ["Interprétation", "Improvisation", "Mémorisation", "Expression", "Travail d'équipe"],
    "Réalisateur": ["Mise en scène", "Direction d'acteurs", "Montage", "Scénario", "Créativité"],
    "Scénariste": ["Écriture", "Développement de personnages", "Structure narrative", "Créativité", "Recherche"],
    "Producteur": ["Gestion de projets", "Financement", "Coordination", "Négociation", "Communication"],
    "Costumier": ["Conception de costumes", "Couture", "Recherche historique", "Créativité", "Communication"],
    "Maquilleur": ["Maquillage", "Effets spéciaux", "Connaissance des produits", "Créativité", "Précision"],
    "Décorateur": ["Conception de décors", "Aménagement", "Créativité", "Budgétisation", "Communication"],
    "Éclairagiste": ["Conception d'éclairage", "Installation", "Programmation", "Créativité", "Technologie"],
    "Technicien son": ["Gestion du son", "Mixage", "Installation", "Dépannage", "Technologie"],
    "Technicien lumière": ["Installation lumière", "Programmation", "Dépannage", "Connaissance des équipements", "Créativité"],
    "Styliste": ["Conception de mode", "Tendances", "Croquis", "Créativité", "Communication"],
    "Modéliste": ["Patronage", "Couture", "Conception", "Précision", "Connaissance des tissus"],
    "Couturier": ["Couture", "Confection", "Précision", "Connaissance des tissus", "Service client"],
    "Bijoutier": ["Conception de bijoux", "Travail des métaux", "Sertissage", "Créativité", "Précision"],
    "Horloger": ["Réparation de montres", "Assemblage", "Précision", "Connaissance mécanique", "Service client"],
    "Ébéniste": ["Travail du bois", "Conception de meubles", "Finition", "Précision", "Créativité"],
    "Menuisier": ["Travail du bois", "Assemblage", "Lecture de plans", "Utilisation d'outils", "Précision"],
    "Charpentier": ["Construction de structures", "Lecture de plans", "Travail du bois", "Normes de sécurité", "Précision"],
    "Maçon": ["Construction", "Lecture de plans", "Maçonnerie", "Normes de sécurité", "Travail d'équipe"],
    "Peintre en bâtiment": ["Peinture", "Préparation de surfaces", "Connaissance des matériaux", "Précision", "Service client"],
    "Carreleur": ["Pose de carrelage", "Coupe", "Préparation de surfaces", "Précision", "Normes de qualité"],
    "Couvreur": ["Pose de toitures", "Réparation", "Normes de sécurité", "Connaissance des matériaux", "Travail en hauteur"],
    "Vitrier": ["Pose de vitres", "Coupe", "Réparation", "Précision", "Service client"],
    "Serrurier": ["Fabrication de serrures", "Installation", "Dépannage", "Connaissance des mécanismes", "Service client"],
    "Forgeron": ["Forge", "Travail des métaux", "Conception", "Précision", "Créativité"],
    "Soudeur": ["Soudure", "Connaissance des métaux", "Normes de sécurité", "Précision", "Utilisation d'équipements"],
    "Métallurgiste": ["Analyse des métaux", "Production", "Connaissance des alliages", "Sécurité", "Rapport"],
    "Conducteur d'engins": ["Conduite", "Maintenance", "Normes de sécurité", "Coordination", "Précision"],
    "Grutier": ["Opération de grues", "Normes de sécurité", "Coordination", "Communication", "Précision"],
    "Chauffeur poids lourd": ["Conduite", "Navigation", "Entretien", "Respect des horaires", "Normes de sécurité"],
    "Livreur": ["Livraison", "Navigation", "Service client", "Gestion du temps", "Conduite sécuritaire"],
    "Magasinier": ["Gestion des stocks", "Organisation", "Réception", "Expédition", "Informatique"],
    "Employé de banque": ["Service client", "Gestion de transactions", "Connaissance bancaire", "Communication", "Précision"],
    "Caissier": ["Gestion de caisse", "Service client", "Comptabilité de base", "Communication", "Travail sous pression"],
    "Vendeur": ["Vente", "Service client", "Négociation", "Connaissance des produits", "Communication"],
    "Manager de magasin": ["Gestion d'équipe", "Vente", "Gestion des stocks", "Budgétisation", "Communication"],
    "Merchandiser": ["Présentation visuelle", "Analyse des ventes", "Gestion des stocks", "Créativité", "Communication"],
    "Sommelier": ["Connaissance des vins", "Dégustation", "Conseil client", "Gestion de cave", "Communication"],
    "Œnologue": ["Vinification", "Analyse chimique", "Dégustation", "Recherche", "Gestion de production"],
    "Viticulteur": ["Culture de la vigne", "Récolte", "Entretien", "Connaissance des sols", "Gestion"],
    "Brasseur": ["Brassage", "Connaissance des ingrédients", "Contrôle qualité", "Hygiène", "Créativité"],
    "Distillateur": ["Distillation", "Connaissance des alcools", "Contrôle qualité", "Hygiène", "Précision"],
    "Apiculteur": ["Gestion des ruches", "Récolte de miel", "Connaissance des abeilles", "Entretien", "Commercialisation"],
    "Éleveur": ["Soins aux animaux", "Gestion d'élevage", "Connaissance animale", "Commercialisation", "Maintenance"],
    "Vacher": ["Traite", "Soins aux vaches", "Entretien", "Connaissance animale", "Hygiène"],
    "Berger": ["Gestion du troupeau", "Soins aux animaux", "Navigation", "Connaissance animale", "Entretien"],
    "Maraîcher": ["Culture maraîchère", "Récolte", "Irrigation", "Connaissance des sols", "Commercialisation"],
    "Arboriculteur": ["Culture d'arbres fruitiers", "Taille", "Récolte", "Connaissance des sols", "Entretien"],
    "Pépiniériste": ["Culture de plantes", "Greffage", "Entretien", "Connaissance botanique", "Commercialisation"],
    "Semencier": ["Production de semences", "Contrôle qualité", "Recherche", "Connaissance botanique", "Commercialisation"],
    "Paysan": ["Agriculture", "Élevage", "Gestion de ferme", "Connaissance des sols", "Commercialisation"],
    "Technicien agricole": ["Analyse des sols", "Gestion des cultures", "Utilisation de machines", "Rapport", "Conseil"],
    "Vulcanologue": ["Analyse volcanique", "Recherche", "Cartographie", "Évaluation des risques", "Rapport"],
    "Sismologue": ["Analyse sismique", "Modélisation", "Recherche", "Évaluation des risques", "Rapport"],
    "Hydrologue": ["Analyse de l'eau", "Modélisation", "Recherche", "Gestion des ressources", "Rapport"],
    "Glaciologue": ["Analyse des glaciers", "Recherche", "Échantillonnage", "Modélisation", "Rapport"],
    "Cartographe": ["Cartographie", "SIG", "Analyse spatiale", "Précision", "Logiciels"],
    "Topographe": ["Mesures topographiques", "Cartographie", "Utilisation d'équipements", "Précision", "Rapport"],
    "Géomaticien": ["SIG", "Analyse spatiale", "Cartographie", "Programmation", "Recherche"],
    "Épidémiologiste": ["Analyse des données", "Recherche", "Prévention", "Rapport", "Statistiques"],
    "Généticien": ["Analyse génétique", "Recherche", "Expérimentation", "Rédaction scientifique", "Bioinformatique"],
    "Biochimiste": ["Analyse biochimique", "Expérimentation", "Recherche", "Rédaction scientifique", "Sécurité"],
    "Microbiologiste": ["Analyse microbiologique", "Expérimentation", "Recherche", "Sécurité", "Rapport"],
    "Pharmacologue": ["Recherche pharmacologique", "Expérimentation", "Analyse", "Sécurité", "Rédaction"],
    "Toxicologue": ["Analyse toxicologique", "Recherche", "Évaluation des risques", "Rapport", "Sécurité"],
    "Nutritionniste": ["Conseil alimentaire", "Évaluation", "Planification", "Communication", "Recherche"],
    "Diététicien": ["Planification alimentaire", "Conseil", "Évaluation", "Communication", "Suivi"],
    "Préparateur physique": ["Entraînement", "Évaluation physique", "Planification", "Motivation", "Communication"],
    "Entraîneur": ["Coaching", "Planification", "Motivation", "Évaluation", "Communication"],
    "Arbitre": ["Connaissance des règles", "Prise de décision", "Communication", "Impartialité", "Gestion du stress"],
    "Agent sportif": ["Négociation", "Gestion de carrière", "Connaissance du sport", "Communication", "Réseautage"],
    "Journaliste sportif": ["Rédaction", "Interview", "Connaissance sportive", "Montage", "Communication"],
    "Commentateur": ["Commentaire en direct", "Connaissance sportive", "Communication", "Improvisation", "Charisme"],
    "Éditeur": ["Révision", "Planification éditoriale", "Coordination", "Connaissance du marché", "Communication"],
    "Libraire": ["Vente", "Conseil client", "Gestion des stocks", "Connaissance littéraire", "Communication"],
    "Imprimeur": ["Impression", "Gestion des machines", "Contrôle qualité", "Organisation", "Précision"],
    "Relieur": ["Reliure", "Restauration", "Précision", "Connaissance des matériaux", "Créativité"]
}

# Descriptions spécifiques par métier
job_descriptions = {
    "Enseignant": [
        "Enseignant avec {years} ans d’expérience, expert en {skill1}, j’ai conçu des cours engageants et utilisé {skill2} pour optimiser l’apprentissage.",
        "Spécialisé en {skill1}, j’ai géré des classes diversifiées pendant {years} ans, maîtrisant {skill2} pour motiver les élèves."
    ],
    "Médecin": [
        "Médecin avec {years} ans d’expérience, j’ai effectué des diagnostics précis en {skill1} et géré des urgences avec {skill2}.",
        "Spécialisé en {skill1}, j’ai travaillé {years} ans en milieu hospitalier, excelling en {skill2} et en soins patients."
    ]
}
for job in jobs:
    if job not in job_descriptions:
        job_descriptions[job] = [
            f"Professionnel(le) avec {{years}} ans d’expérience en {{skill1}}, compétent(e) en {{skill2}} et motivé(e) à relever des défis.",
            f"Expert(e) en {{skill1}} avec {{years}} ans d’expérience, je maîtrise {{skill2}} pour des résultats optimaux."
        ]

# Langues parlées avec probabilités
nationality_languages = {
    "Djibouti": {"Arabe": 0.9, "Français": 0.8, "Somali": 0.5, "Afar": 0.5, "Anglais": 0.2},
    "French": {"Français": 0.95, "Anglais": 0.5, "Arabe": 0.1},
    "English": {"Anglais": 0.95, "Français": 0.2},
    "Russian": {"Russe": 0.95, "Anglais": 0.4},
    "Italian": {"Italien": 0.95, "Anglais": 0.5},
    "Ethiopian": {"Amharique": 0.9, "Anglais": 0.4},
    "Other": {"Arabe": 0.7, "Anglais": 0.6, "Français": 0.3}
}

# Niveaux de maîtrise
language_proficiency = ["Courant", "Intermédiaire", "Débutant"]

# Adresses djiboutiennes
djiboutian_addresses = [
    "Rue de Balbala, Djibouti", "Quartier Ambouli, Djibouti", "Avenue Gabode, Djibouti",
    "Cité Arhiba, Djibouti", "Boulevard de la République, Djibouti"
]

# Institutions djiboutiennes
institutions = [
    "Université de Djibouti", "Lycée d’État de Djibouti", "Institut de Formation Professionnelle",
    "Centre de Formation Technique de Balbala", "École Nationale de Commerce"
]

# Entreprises djiboutiennes
djiboutian_companies = [
    "Port de Djibouti", "Ministère de l’Éducation", "Djibouti Telecom", "Hôpital Général Peltier",
    "Banque de Djibouti", "Électricité de Djibouti", "Aéroport International de Djibouti"
]

# Certificats
certificates = [
    "CAPES", "AWS Certified Solutions Architect", "Cisco CCNA", "PMP", "Scrum Master",
    "Google Data Analytics", "Coursera Machine Learning", "IELTS", "TOEFL", "CFA Level 1",
    "Microsoft Certified: Azure Fundamentals", "CompTIA Security+", "First Aid Certification"
]

# Fonction pour assigner les langues
def assign_languages(nationality, is_djiboutian):
    langs = []
    lang_probs = nationality_languages["Djibouti" if is_djiboutian else nationality]
    for lang, prob in lang_probs.items():
        if random.random() < prob:
            proficiency = random.choice(language_proficiency)
            langs.append(f"{lang} ({proficiency})")
    return langs[:random.randint(2, 4)] or ["Français (Courant)"]

# Fonction pour générer des expériences professionnelles
def generate_experience(job, years, profile_type):
    num_experiences = {"junior": random.randint(1, 2), "intermediate": random.randint(2, 4),
                       "experienced": random.randint(3, 5), "senior": random.randint(4, 6)}[profile_type]
    experiences = []
    current_year = 2025
    available_years = years
    for _ in range(num_experiences):
        if available_years < 1:  # If no more years to assign, break
            break
        
        company = random.choice(djiboutian_companies) if random.random() < 0.6 else fake_fr.company()
        
        # Ensure duration is at least 1 and not more than available_years or 5.
        # upper_bound_duration will be at least 1 due to the check above.
        upper_bound_duration = min(5, available_years)
        duration = random.randint(1, upper_bound_duration)
        
        start_year = current_year - duration
        role = f"{job} {random.choice(['Junior', 'Senior', ''])}".strip()
        responsibilities = random.sample(job_skills[job], random.randint(2, min(5, len(job_skills[job]))))
        experiences.append(f"{role} chez {company} ({start_year}-{current_year}): {', '.join(responsibilities)}")
        current_year -= duration
        available_years -= duration
    return experiences

# Fonction pour générer une description
def generate_description(job, years, skills, language):
    template = random.choice(job_descriptions[job])
    description = template.format(years=years, skill1=skills[0], skill2=skills[1])
    return description if language == "Français" else translator.translate(description, dest="en").text

# Fonction pour générer une référence
def generate_reference():
    name = fake_fr.name() if random.random() < 0.7 else fake_en.name()
    title = random.choice(["Manager", "Collègue", "Superviseur"])
    email = f"{name.lower().replace(' ', '.')}"
    return f"{name}, {title}, {email}"

# Fonction pour valider le CV
def validate_cv(cv):
    profile_limits = {"junior": 5, "intermediate": 10, "experienced": 15, "senior": 20}
    if cv["Années d’expérience"] > profile_limits[cv['Profil']]:
        cv["Années d’expérience"] = random.randint(0, profile_limits[cv['Profil']])
    if cv["Métier"] not in job_skills:
        raise ValueError(f"Compétences manquantes pour {cv['Métier']}")

# Fonction pour vérifier l'unicité du CV
cv_signatures = set()
def is_unique_cv(cv):
    signature = f"{cv['Nom complet']}_{cv['Métier']}_{cv['Formation']}_{cv['Institution']}"
    if signature in cv_signatures:
        return False
    cv_signatures.add(signature)
    return True

# Fonction pour générer un CV
def generate_cv(job, profile_type, is_djiboutian, nationality=None, cv_index=0):
    if is_djiboutian:
        first_name = random.choice(djiboutian_first_names)
        middle_name = random.choice(djiboutian_first_names)
        last_name = random.choice(djiboutian_last_names)
        full_name = f"{first_name} {middle_name} {last_name}"
        nationality = "Djibouti"
        address = random.choice(djiboutian_addresses)
    else:
        nationality = nationality or random.choice(["French", "English", "Russian", "Italian", "Ethiopian", "Other"])
        first_name = random.choice(foreign_first_names[nationality])
        last_name = random.choice(foreign_last_names)
        full_name = f"{first_name} {last_name}"
        address = fake_fr.address() if nationality == "French" else fake_en.address()
    
    # Expérience
    experience_years = {"junior": random.randint(0, 5), "intermediate": random.randint(6, 10),
                        "experienced": random.randint(11, 15), "senior": random.randint(16, 20)}[profile_type]
    
    # Compétences
    num_skills = random.randint(4, min(8, len(job_skills[job])))
    skills = random.sample(job_skills[job], num_skills)
    
    # Formation
    education = random.choice([
        "Baccalauréat", f"Licence en {job}", f"Master en {job}", "Diplôme professionnel",
        f"BTS en {job}", "Doctorat en {job}"
    ])
    institution = random.choice(institutions)
    
    # Certificats
    num_certs = random.randint(1, 3)
    certificate_details = random.sample(certificates, num_certs)
    certificate_details = [f"{cert} ({random.choice(institutions)})" for cert in certificate_details]
    
    # Langues
    spoken_languages = assign_languages(nationality, is_djiboutian)
    
    # Hobbies
    hobbies = random.sample(["Football", "Lecture", "Cuisine", "Voyage", "Photographie", "Musique"], random.randint(2, 4))
    
    # Langue du CV
    cv_language = "Français" if random.random() < 0.9 else "Anglais"
    
    # Description
    description = generate_description(job, experience_years, skills, cv_language)
    
    # Expériences
    experiences = generate_experience(job, experience_years, profile_type)
    
    # Références
    references = [generate_reference() for _ in range(random.randint(1, 2))]
    
    # Email et téléphone
    email = f"{first_name.lower()}.{last_name.lower()}@example.com".replace("'", "").replace(" ", "")
    phone = f"+25377{random.randint(100000, 999999)}"
    
    # Format
    # Adjust format determination based on the new max_cvs for a 50/50 split
    cv_format = "pdf" if cv_index < (max_cvs / 2) else "docx"
    
    # Template de mise en page
    template = random.choice(["classic", "modern", "compact"])
    
    cv = {
        "Nom complet": full_name,
        "Nationalité": nationality,
        "Adresse": address,
        "Métier": job,
        "Années d’expérience": experience_years,
        "Profil": profile_type,
        "Formation": education,
        "Institution": institution,
        "Compétences": ", ".join(skills),
        "Certificats": ", ".join(certificate_details),
        "Langues parlées": ", ".join(spoken_languages),
        "Hobbies": ", ".join(hobbies),
        "Description": description,
        "Expériences": experiences,
        "Références": references,
        "Email": email,
        "Téléphone": phone,
        "Langue du CV": cv_language,
        "Format": cv_format,
        "Template": template
    }
    
    validate_cv(cv)
    return cv

# Fonction pour créer un PDF
def create_cv_pdf(cv, output_path):
    try:
        doc = SimpleDocTemplate(output_path, pagesize=letter, title=f"CV {cv['Nom complet']}")
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(name='Title', fontName='Helvetica-Bold', fontSize=16, alignment=1, spaceAfter=12)
        heading_style = ParagraphStyle(name='Heading', fontName='Helvetica-Bold', fontSize=12, spaceAfter=8)
        normal_style = ParagraphStyle(name='Normal', fontName='Helvetica', fontSize=10, spaceAfter=6)
        
        content = []
        content.append(Paragraph(cv["Nom complet"], title_style))
        content.append(Paragraph(f"{cv['Email']} | {cv['Téléphone']}", normal_style))
        content.append(Spacer(1, 0.1 * inch))
        
        if cv["Template"] == "modern":
            data = [[Paragraph("Compétences: " + cv["Compétences"], normal_style),
                     Paragraph("Expériences: " + "<br/>".join(cv["Expériences"]), normal_style)]]
            table = Table(data, colWidths=[3 * inch, 3 * inch])
            table.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP')]))
            content.append(table)
        else:
            content.append(Paragraph("Informations personnelles" if cv["Langue du CV"] == "Français" else "Personal Information", heading_style))
            content.append(Paragraph(f"Nationalité: {cv['Nationalité']}<br/>Adresse: {cv['Adresse']}", normal_style))
            content.append(Spacer(1, 8))
            content.append(Paragraph("Expériences professionnelles" if cv["Langue du CV"] == "Français" else "Professional Experience", heading_style))
            for exp in cv["Expériences"]:
                content.append(Paragraph(exp, normal_style))
            content.append(Paragraph("Compétences" if cv["Langue du CV"] == "Français" else "Skills", heading_style))
            content.append(Paragraph(cv["Compétences"], normal_style))
        
        content.append(Paragraph("Formation" if cv["Langue du CV"] == "Français" else "Education", heading_style))
        content.append(Paragraph(f"{cv['Formation']} ({cv['Institution']})", normal_style))
        content.append(Paragraph("Certificats" if cv["Langue du CV"] == "Français" else "Certificates", heading_style))
        content.append(Paragraph(cv["Certificats"], normal_style))
        content.append(Paragraph("Langues parlées" if cv["Langue du CV"] == "Français" else "Spoken Languages", heading_style))
        content.append(Paragraph(cv["Langues parlées"], normal_style))
        content.append(Paragraph("Hobbies" if cv["Langue du CV"] == "Français" else "Hobbies", heading_style))
        content.append(Paragraph(cv["Hobbies"], normal_style))
        content.append(Paragraph("Description" if cv["Langue du CV"] == "Français" else "Profile Description", heading_style))
        content.append(Paragraph(cv["Description"], normal_style))
        content.append(Paragraph("Références" if cv["Langue du CV"] == "Français" else "References", heading_style))
        for ref in cv["Références"]:
            content.append(Paragraph(ref, normal_style))
        
        doc.build(content)
    except Exception as e:
        print(f"Erreur lors de la création du PDF {output_path}: {e}")

# Fonction pour créer un DOCX
def create_cv_docx(cv, output_path):
    try:
        doc = Document()
        title = doc.add_paragraph(cv["Nom complet"])
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.runs[0].font.size = Pt(16)
        title.runs[0].font.bold = True
        
        doc.add_paragraph(f"{cv['Email']} | {cv['Téléphone']}")
        
        if cv["Template"] == "modern":
            table = doc.add_table(rows=1, cols=2)
            table.cells[0, 0].text = "Compétences\n" + cv["Compétences"]
            table.cells[0, 1].text = "Expériences\n" + "\n".join(cv["Expériences"])
        else:
            doc.add_paragraph("Informations personnelles" if cv["Langue du CV"] == "Français" else "Personal Information").runs[0].font.size = Pt(12)
            doc.add_paragraph(f"Nationalité: {cv['Nationalité']}\nAdresse: {cv['Adresse']}")
            doc.add_paragraph("Expériences professionnelles" if cv["Langue du CV"] == "Français" else "Professional Experience").runs[0].font.size = Pt(12)
            for exp in cv["Expériences"]:
                doc.add_paragraph(exp)
            doc.add_paragraph("Compétences" if cv["Langue du CV"] == "Français" else "Skills").runs[0].font.size = Pt(12)
            doc.add_paragraph(cv["Compétences"])
        
        doc.add_paragraph("Formation" if cv["Langue du CV"] == "Français" else "Education").runs[0].font.size = Pt(12)
        doc.add_paragraph(f"{cv['Formation']} ({cv['Institution']})")
        doc.add_paragraph("Certificats" if cv["Langue du CV"] == "Français" else "Certificates").runs[0].font.size = Pt(12)
        doc.add_paragraph(cv["Certificats"])
        doc.add_paragraph("Langues parlées" if cv["Langue du CV"] == "Français" else "Spoken Languages").runs[0].font.size = Pt(12)
        doc.add_paragraph(cv["Langues parlées"])
        doc.add_paragraph("Hobbies" if cv["Langue du CV"] == "Français" else "Hobbies").runs[0].font.size = Pt(12)
        doc.add_paragraph(cv["Hobbies"])
        doc.add_paragraph("Description" if cv["Langue du CV"] == "Français" else "Profile Description").runs[0].font.size = Pt(12)
        doc.add_paragraph(cv["Description"])
        doc.add_paragraph("Références" if cv["Langue du CV"] == "Français" else "References").runs[0].font.size = Pt(12)
        for ref in cv["Références"]:
            doc.add_paragraph(ref)
        
        doc.save(output_path)
    except Exception as e:
        print(f"Erreur lors de la création du DOCX {output_path}: {e}")

# Fonction pour générer et sauvegarder un CV
def generate_and_save_cv(args):
    cv, index = args
    try:
        filename = f"cv_{cv['Nom complet'].replace(' ', '_')}_{cv['Métier'].replace(' ', '_')}_{random.randint(1000,9999)}.{cv['Format']}"
        # Define base output directory
        base_output_dir = "generated_cvs_output"
        output_path = os.path.join(base_output_dir, f"cvs_{cv['Format']}", filename)
        
        # Ensure the specific format directory exists
        os.makedirs(os.path.join(base_output_dir, f"cvs_{cv['Format']}"), exist_ok=True)

        if cv["Format"] == "pdf":
            create_cv_pdf(cv, output_path)
        else:
            create_cv_docx(cv, output_path)
        return {"filename": filename, "cv": cv}
    except Exception as e:
        print(f"Erreur lors de la génération du CV {index}: {e}")
        return None

# Étape 4 : Générer 14 400 CV uniques (72 CV par métier)
cvs = []
# Define base output directory and create it
base_output_dir = "generated_cvs_output"
os.makedirs(base_output_dir, exist_ok=True)
os.makedirs(os.path.join(base_output_dir, 'cvs_pdf'), exist_ok=True)
os.makedirs(os.path.join(base_output_dir, 'cvs_docx'), exist_ok=True)

# Répartition des profils
profile_types = ["junior", "intermediate", "experienced", "senior"]
cv_per_job = 1  # Generate 1 CV per job to reach 200 CVs for 200 jobs
total_cvs = 0
max_cvs = 200 # Generate only 200 files

# Générer les CVs
with ThreadPoolExecutor(max_workers=8) as executor:
    args = []
    for job in jobs:
        cv_count = 0
        attempts = 0
        max_attempts = 100  # Prevent infinite loops
        while cv_count < cv_per_job and attempts < max_attempts:
            is_djiboutian = random.random() < 0.85
            profile_type = random.choice(profile_types)
            cv = generate_cv(job, profile_type, is_djiboutian, cv_index=total_cvs)
            if is_unique_cv(cv):
                args.append((cv, total_cvs))
                cv_count += 1
                total_cvs += 1
            attempts += 1
            if total_cvs >= max_cvs:
                break
        if total_cvs >= max_cvs:
            break
    
    results = list(executor.map(generate_and_save_cv, args))
    cvs = [result["cv"] for result in results if result]

# Étape 5 : Exporter les métadonnées
cv_metadata = pd.DataFrame(cvs)
metadata_path = os.path.join(base_output_dir, "cv_metadata.csv")
cv_metadata.to_csv(metadata_path, index=False)
print(f"Métadonnées exportées dans {metadata_path}")

# Étape 6 : Créer une archive ZIP
zip_filename = os.path.join(base_output_dir, f'{max_cvs}_cvs.zip') # Dynamic zip filename
with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED, compresslevel=9) as zipf:
    pdf_dir_to_zip = os.path.join(base_output_dir, 'cvs_pdf')
    for root, _, files in os.walk(pdf_dir_to_zip):
        for file in files:
            if file.endswith('.pdf'):
                full_path = os.path.join(root, file)
                # Add file to zip with relative path inside 'cvs_pdf' folder in zip
                zipf.write(full_path, os.path.join('cvs_pdf', os.path.relpath(full_path, pdf_dir_to_zip)))

    docx_dir_to_zip = os.path.join(base_output_dir, 'cvs_docx')
    for root, _, files in os.walk(docx_dir_to_zip):
        for file in files:
            if file.endswith('.docx'):
                full_path = os.path.join(root, file)
                # Add file to zip with relative path inside 'cvs_docx' folder in zip
                zipf.write(full_path, os.path.join('cvs_docx', os.path.relpath(full_path, docx_dir_to_zip)))
    
    zipf.write(metadata_path, os.path.basename(metadata_path))

print(f"Archive ZIP créée : {zip_filename}")

# Étape 7 : Télécharger l'archive
# from google.colab import files
# files.download(zip_filename)
print(f"L'archive ZIP est disponible localement à : {os.path.abspath(zip_filename)}")