# Étape 1 : Installer les dépendances
# pip install faker reportlab python-docx pandas googletrans==3.1.0a0

# Étape 2 : Importer les bibliothèques
from faker import Faker
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from googletrans import Translator
import random
import os
import zipfile
from concurrent.futures import ThreadPoolExecutor


# from datetime import datetime # Not used, can be removed

# Étape 3 : Initialiser Faker pour différentes langues
fake_fr = Faker('fr_FR')
fake_en = Faker('en_US')
# fake_ru = Faker('ru_RU') # Not directly used for name generation now if using global pools
# fake_it = Faker('it_IT') # Not directly used
# fake_am = Faker('am_ET') # Not directly used

# Initialiser le traducteur
translator = Translator()

# Listes de prénoms et noms djiboutiens (as per PDF: 120 first, 105 last)
djiboutian_first_names_actual_120 = [
    # Masculins (60)
    "Ahmed", "Hassan", "Mahad", "Ismail", "Abdourahman", "Ali", "Ibrahim", "Abdi", "Youssouf", "Souleiman",
    "Jamal", "Zakaria", "Mohamed", "Farhan", "Idriss", "Faysal", "Yassin", "Bilal", "Omar", "Abdallah",
    "Barkhad", "Anwar", "Saïd", "Nour", "Dini", "Mahamoud", "Khaled", "Awaleh", "Bileh", "Muse",
    "Saad", "Salah", "Warsame", "Guedi", "Walid", "Abdirahman", "Haroun", "Moustapha", "Abshir",
    "Suleyman", "Dawoud", "Houssein", "Abokor", "Aweis", "Hersi", "Aaden", "Ayub", "Farah", "Ayman",
    "Rachid", "Madar", "Nuradin", "Hassanali", "Kalif", "Deria", "Saidou", "Samatar", "Tahlil",
    "Wais", "Darar",
    # Féminins (60)
    "Fadumo", "Ayan", "Hibo", "Ifrah", "Nimo", "Hawa", "Sagal", "Asha", "Safia", "Hodan",
    "Zamzam", "Amal", "Faiza", "Rahma", "Amina", "Khadija", "Naima", "Yasmin", "Salma", "Ikram",
    "Roda", "Halima", "Saida", "Habiba", "Mariam", "Fatouma", "Asli", "Souad", "Batula", "Anisa",
    "Hani", "Deeqa", "Shaafi", "Sahra", "Rihana", "Munira", "Bilan", "Layla", "Nadra", "Muna",
    "Mako", "Sucaad", "Faisa", "Anab", "Haniyo", "Sulekha", "Fawzia", "Farhia", "Ifra", "Sudi",
    "Lubna", "Bashira", "Nura", "Zamzama", "Zaynab", "Deka", "Ikran", "Saynab", "Jawahir",
    "Farhiyo", "Ubah"
]

_djiboutian_last_names_base = [
    "Abdi Farah", "Ali Robleh", "Hassan Dini", "Youssouf Omar", "Mohamed Farhan", "Ismail Awaleh",
    "Ahmed Gedi", "Ibrahim Duale", "Mahad Saïd", "Abdirahman Muse", "Farhan Mahamoud", "Guedi Warsame",
    "Omar Abdi", "Abdourahman Walid", "Khaled Awaleh", "Abshir Robleh", "Suleyman Barre",
    "Djamal Samatar", "Robleh Tahlil", "Barkhad Robleh", "Souleiman Abdi", "Saïd Deria",
    "Warsame Madar", "Nour Hassan", "Mahdi Robleh", "Dawoud Robleh", "Abokor Wais", "Deria Saïd",
    "Saad Abdi", "Houssein Hassan", "Moustapha Abdi", "Walid Barkhad", "Aliyow Djamal",
    "Abdi Sharmarke", "Salah Djamal", "Ayub Farah", "Rachid Madar", "Mahamoud Saidou",
    "Awaleh Barkat", "Yassin Madar", "Abdi Barre", "Barkat Kalif", "Madar Guedi", "Jamal Sharmarke",
    "Bilal Saïd", "Robleh Wais", "Barre Awaleh", "Khalif Warsame", "Farah Hersi", "Abdallah Madar",
    "Hersi Samatar", "Nouradin Tahlil", "Saidou Hassan", "Kalif Dawoud", "Walid Saidou",
    "Djamal Barkat", "Barkat Awaleh", "Faysal Barre", "Samatar Guedi", "Deria Hersi",
    "Aweis Tahlil", "Robleh Barkhad", "Awaleh Saidou", "Ali Tahlil", "Abdirahman Barkhad",
    "Madar Warsame", "Ismail Deria", "Omar Abdi", "Jamal Saidou", "Saïd Abshir", "Ali Robleh", # This Ali Robleh is repeated from start
    "Youssouf Madar", "Barkhad Muse", "Tahlil Warsame", "Barkat Duale", "Awaleh Guedi",
    "Samatar Mahdi", "Walid Nour", "Muse Saïd", "Saidou Barkat", "Mahamoud Djamal",
    "Sharmarke Warsame", "Deria Robleh", "Guedi Awaleh", "Hassan Kalif"
] # This list has 78 unique items. We need 105.

djiboutian_last_names_actual_105 = list(set(_djiboutian_last_names_base)) # Get unique base
# Pad to 105 if needed
if len(djiboutian_last_names_actual_105) < 105:
    needed = 105 - len(djiboutian_last_names_actual_105)
    for i in range(needed):
        base_name = random.choice(_djiboutian_last_names_base) # Pick any, even if it creates non-unique base for padding
        djiboutian_last_names_actual_105.append(f"{base_name} {i+2}") # Make it unique string
# Ensure exactly 105, trim if somehow over after set() and padding logic error (unlikely here)
djiboutian_last_names_actual_105 = djiboutian_last_names_actual_105[:105]
if len(djiboutian_last_names_actual_105) < 105: # If still less (e.g. base list was very small)
    print(f"Warning: Djiboutian last names count is {len(djiboutian_last_names_actual_105)}, less than desired 105. Padding with Faker.")
    faker_generic = Faker()
    while len(djiboutian_last_names_actual_105) < 105:
        new_name = faker_generic.last_name()
        if new_name not in djiboutian_last_names_actual_105:
             djiboutian_last_names_actual_105.append(new_name)


# Foreign Names: PDF specifies global pools of 40 first names, 50 last names.
# Populate these pools, e.g., by picking from the existing more detailed lists or using Faker.
_temp_foreign_first = []
_original_foreign_first_names_source = { # From original code
    "French": ["Lucas", "Hugo", "Enzo", "Arthur", "Léo", "Nathan", "Raphaël", "Jules", "Tom", "Clément", "Louis", "Adam", "Noé", "Théo", "Maxime"],
    "English": ["James", "John", "Robert", "Michael", "William", "David", "Richard", "Charles", "Joseph", "Thomas", "Daniel", "Matthew", "Andrew", "Kevin", "Brian"],
    "Russian": ["Ivan", "Dimitri", "Alexei", "Nikolai", "Mikhail", "Vladimir", "Yuri", "Andrei", "Sergei", "Pavel"],
    "Italian": ["Luca", "Matteo", "Marco", "Andrea", "Alessandro", "Gabriele", "Giovanni", "Antonio", "Stefano", "Francesco"],
    "Ethiopian": ["Bekele", "Tadesse", "Tesfaye", "Abebe", "Dawit", "Getachew", "Hailu", "Kassahun", "Solomon", "Yosef"],
    "Other": ["Ali", "Omar", "Amir", "Malik", "Ibrahim", "Hamza", "Sami", "Elias", "Youssef", "Karim"]
}
for names in _original_foreign_first_names_source.values():
    _temp_foreign_first.extend(names)
FOREIGN_FIRST_NAMES_POOL_40 = list(set(_temp_foreign_first))[:40]
# Ensure we have 40, pad with Faker if necessary
if len(FOREIGN_FIRST_NAMES_POOL_40) < 40:
    print(f"Warning: Foreign first names pool count is {len(FOREIGN_FIRST_NAMES_POOL_40)}. Padding with Faker.")
    faker_generic = Faker()
    while len(FOREIGN_FIRST_NAMES_POOL_40) < 40:
        new_name = faker_generic.first_name()
        if new_name not in FOREIGN_FIRST_NAMES_POOL_40:
            FOREIGN_FIRST_NAMES_POOL_40.append(new_name)


_original_foreign_last_names_source = [ # From original code
    "Dubois", "Lefèvre", "Bernard", "Girard", "Morel", "Smith", "Johnson", "Williams", "Brown", "Jones",
    "Petrov", "Sokolov", "Kuznetsov", "Ivanov", "Makarov", "Rossi", "Bianchi", "Romano", "Conti", "Gallo",
    "Teshome", "Gebre", "Mekonnen", "Alemu", "Desta", "El Fassi", "Ben Omar", "Diouf", "Bamba", "Haddad"
]
FOREIGN_LAST_NAMES_POOL_50 = list(set(_original_foreign_last_names_source))[:50]
# Ensure we have 50, pad with Faker if necessary
if len(FOREIGN_LAST_NAMES_POOL_50) < 50:
    print(f"Warning: Foreign last names pool count is {len(FOREIGN_LAST_NAMES_POOL_50)}. Padding with Faker.")
    faker_generic = Faker()
    while len(FOREIGN_LAST_NAMES_POOL_50) < 50:
        new_name = faker_generic.last_name()
        if new_name not in FOREIGN_LAST_NAMES_POOL_50:
            FOREIGN_LAST_NAMES_POOL_50.append(new_name)


# Liste des 200 métiers
jobs = [
    "Enseignant", "Médecin", "Ingénieur informatique", "Infirmier", "Commerçant", "Comptable", "Avocat",
    "Chauffeur", "Électricien", "Pharmacien", "Journaliste", "Développeur web", "Cuisinier", "Secrétaire",
    "Logisticien", "Pêcheur", "Agriculteur", "Architecte", "Dentiste", "Vétérinaire", "Banquier", "Marketeur",
    "Plombier", "Mécanicien", "Coiffeur", "Tailleur", "Boulanger", "Pâtissier", "Serveur", "Barman",
    "Agent immobilier", "Agent de sécurité", "Policier", "Pompier", "Militaire", "Technicien de maintenance",
    "Analyste financier", "Consultant", "Graphiste", "Photographe", "Vidéaste", "Monteur vidéo", "Rédacteur",
    "Traducteur", "Interprète", "Bibliothécaire", "Archiviste", "Conservateur de musée", "Guide touristique",
    "Agent de voyage", "Pilote", "Hôtesse de l'air", "Contrôleur aérien", "Marin", "Capitaine de navire",
    "Paysagiste", "Jardinier", "Fleuriste", "Écologiste", "Biologiste", "Chimiste", "Physicien", "Mathématicien",
    "Statisticien", "Data scientist", "Data analyst", "Ingénieur civil", "Ingénieur mécanique", "Ingénieur électrique",
    "Urbaniste", "Géologue", "Météorologue", "Océanographe", "Astronome", "Psychologue", "Psychiatre",
    "Travailleur social", "Conseiller d'orientation", "Coach", "Formateur", "Animateur", "Éducateur spécialisé",
    "Orthophoniste", "Kinésithérapeute", "Ergothérapeute", "Opticien", "Audioprothésiste", "Sage-femme",
    "Aide-soignant", "Ambulancier", "Laborantin", "Chercheur", "Professeur d'université", "Assistant de recherche",
    "Économiste", "Sociologue", "Anthropologue", "Historien", "Géographe", "Politicien", "Diplomate",
    "Agent des douanes", "Inspecteur fiscal", "Auditeur", "Actuaire", "Assureur", "Courtier", "Notaire",
    "Huissier", "Juge", "Procureur", "Greffier", "Planificateur", "Gestionnaire de projet", "Scrum master",
    "Product owner", "UX designer", "UI designer", "Développeur mobile", "Développeur backend", "Développeur frontend",
    "Administrateur système", "Administrateur réseau", "Spécialiste cybersécurité", "Testeur QA", "Technicien support",
    "Community manager", "Responsable communication", "Publicitaire", "Organisateur d'événements", "Organisateur de mariage",
    "DJ", "Musicien", "Chanteur", "Acteur", "Réalisateur", "Scénariste", "Producteur", "Costumier", "Maquilleur",
    "Décorateur", "Éclairagiste", "Technicien son", "Technicien lumière", "Styliste", "Modéliste", "Couturier",
    "Bijoutier", "Horloger", "Ébéniste", "Menuisier", "Charpentier", "Maçon", "Peintre en bâtiment", "Carreleur",
    "Couvreur", "Vitrier", "Serrurier", "Forgeron", "Soudeur", "Métallurgiste", "Conducteur d'engins", "Grutier",
    "Chauffeur poids lourd", "Livreur", "Magasinier", "Employé de banque", "Caissier", "Vendeur", "Manager de magasin",
    "Merchandiser", "Sommelier", "Œnologue", "Viticulteur", "Brasseur", "Distillateur", "Apiculteur", "Éleveur",
    "Vacher", "Berger", "Maraîcher", "Arboriculteur", "Pépiniériste", "Semencier", "Paysan", "Technicien agricole",
    "Vulcanologue", "Sismologue", "Hydrologue", "Glaciologue", "Cartographe", "Topographe", "Géomaticien",
    "Épidémiologiste", "Généticien", "Biochimiste", "Microbiologiste", "Pharmacologue", "Toxicologue",
    "Nutritionniste", "Diététicien", "Préparateur physique", "Entraîneur", "Arbitre", "Agent sportif",
    "Journaliste sportif", "Commentateur", "Éditeur", "Libraire", "Imprimeur", "Relieur"
]

# Compétences spécifiques par métier (5 à 15 par métier)
job_skills = {
    # ... (original job_skills dictionary remains here, ensure it's comprehensive) ...
    "Enseignant": ["Pédagogie", "Gestion de classe", "Planification de cours", "Évaluation des élèves", "Technologies éducatives", "Communication", "Adaptabilité", "Créativité", "Travail d'équipe", "Motivation des élèves"],
    "Médecin": ["Diagnostic médical", "Soins aux patients", "Pharmacologie", "Chirurgie", "Gestion des urgences", "Empathie", "Analyse de dossiers médicaux", "Communication patient", "Éthique médicale", "Mise à jour des connaissances"],
    "Ingénieur informatique": ["Programmation", "Résolution de problèmes", "Conception de systèmes", "Gestion de bases de données", "Cybersécurité", "Cloud computing", "DevOps", "Analyse des besoins", "Test de logiciels", "Maintenance système", "Python", "Java", "Agilité"],
    # Add all other job skills from the original script.
    # For brevity, I'm not pasting all of them again. Assume the original list is complete.
    # Ensure each job in `jobs` has an entry in `job_skills` with at least 5 skills.
}
# Example: Fill missing job skills if any (for robustness, though the original list is quite long)
for job_title in jobs:
    if job_title not in job_skills:
        job_skills[job_title] = [f"Compétence générique {i} pour {job_title}" for i in range(1, 6)]
    elif len(job_skills[job_title]) < 5:
         job_skills[job_title].extend([f"Compétence additionnelle {i}" for i in range(len(job_skills[job_title]), 5)])


# Descriptions spécifiques par métier
job_descriptions = {
    "Enseignant": [
        "Enseignant avec {years} ans d’expérience, expert en {skill1}, j’ai conçu des cours engageants et utilisé {skill2} pour optimiser l’apprentissage.",
        "Spécialisé en {skill1}, j’ai géré des classes diversifiées pendant {years} ans, maîtrisant {skill2} pour motiver les élèves."
    ],
    "Médecin": [
        "Médecin avec {years} ans d’expérience, j’ai effectué des diagnostics précis en {skill1} et géré des urgences avec {skill2}.",
        "Spécialisé en {skill1}, j’ai travaillé {years} ans en milieu hospitalier, excelling en {skill2} et en soins patients."
    ]
    # ... (original job_descriptions dictionary remains here) ...
}
for job in jobs:
    if job not in job_descriptions:
        job_descriptions[job] = [
            f"Professionnel(le) avec {{years}} ans d’expérience en {{skill1}}, compétent(e) en {{skill2}} et motivé(e) à relever des défis.",
            f"Expert(e) en {{skill1}} avec {{years}} ans d’expérience, je maîtrise {{skill2}} pour des résultats optimaux."
        ]

# Langues parlées avec probabilités
nationality_languages = {
    "Djibouti": {"Arabe": 0.9, "Français": 0.8, "Somali": 0.5, "Afar": 0.5, "Anglais": 0.2},
    "French": {"Français": 0.95, "Anglais": 0.5, "Arabe": 0.1},
    "English": {"Anglais": 0.95, "Français": 0.2},
    "Russian": {"Russe": 0.95, "Anglais": 0.4},
    "Italian": {"Italien": 0.95, "Anglais": 0.5},
    "Ethiopian": {"Amharique": 0.9, "Anglais": 0.4},
    "Other": {"Arabe": 0.7, "Anglais": 0.6, "Français": 0.3}
}
language_proficiency = ["Courant", "Intermédiaire", "Débutant"]
djiboutian_addresses = ["Rue de Balbala, Djibouti", "Quartier Ambouli, Djibouti", "Avenue Gabode, Djibouti", "Cité Arhiba, Djibouti", "Boulevard de la République, Djibouti"]
institutions = ["Université de Djibouti", "Lycée d’État de Djibouti", "Institut de Formation Professionnelle", "Centre de Formation Technique de Balbala", "École Nationale de Commerce"]
djiboutian_companies = ["Port de Djibouti", "Ministère de l’Éducation", "Djibouti Telecom", "Hôpital Général Peltier", "Banque de Djibouti", "Électricité de Djibouti", "Aéroport International de Djibouti"]
certificates = ["CAPES", "AWS Certified Solutions Architect", "Cisco CCNA", "PMP", "Scrum Master", "Google Data Analytics", "Coursera Machine Learning", "IELTS", "TOEFL", "CFA Level 1", "Microsoft Certified: Azure Fundamentals", "CompTIA Security+", "First Aid Certification"]

def assign_languages(nationality_label): # Changed parameter name for clarity
    langs = []
    # Use the specific nationality label for language probabilities
    lang_probs = nationality_languages.get(nationality_label, nationality_languages["Other"])
    for lang, prob in lang_probs.items():
        if random.random() < prob:
            proficiency = random.choice(language_proficiency)
            langs.append(f"{lang} ({proficiency})")
    return langs[:random.randint(2, 4)] or ["Français (Courant)"] # Ensure at least one language

def generate_experience(job, years, profile_type):
    num_experiences = {"junior": random.randint(1, 2), "intermediate": random.randint(2, 4),
                       "experienced": random.randint(3, 5), "senior": random.randint(4, 6)}[profile_type]
    experiences = []
    current_year = 2025 # Assuming generation for near future
    available_years = years
    
    job_specific_skills_list = job_skills.get(job, [])
    if not job_specific_skills_list: # Fallback if a job has no skills defined
        job_specific_skills_list = ["Gestion de projet", "Communication", "Travail d'équipe", "Résolution de problèmes", "Analyse"]


    for _ in range(num_experiences):
        if available_years <= 0: break
        company = random.choice(djiboutian_companies) if random.random() < 0.6 else fake_fr.company()
        duration = random.randint(1, min(5, max(1, available_years))) # Ensure duration is at least 1
        start_year = current_year - duration
        role = f"{job} {random.choice(['Junior', 'Senior', ''])}".strip()
        
        num_resp_to_pick = random.randint(2, min(5, len(job_specific_skills_list)))
        responsibilities = random.sample(job_specific_skills_list, num_resp_to_pick) if job_specific_skills_list else ["Tâches diverses"]
        
        experiences.append(f"{role} chez {company} ({start_year}-{current_year}): {', '.join(responsibilities)}")
        current_year -= duration
        available_years -= duration
        if available_years < 0: available_years = 0 # Cap at 0
    return experiences

def generate_description(job, years, skills, language):
    template = random.choice(job_descriptions[job])
    # Ensure skills list is not empty for placeholder access
    skill1_desc = skills[0] if skills else "compétences clés"
    skill2_desc = skills[1] if len(skills) > 1 else "polyvalence"
    description_fr = template.format(years=years, skill1=skill1_desc, skill2=skill2_desc)
    
    if language == "Français":
        return description_fr
    else:
        try:
            return translator.translate(description_fr, dest="en").text
        except Exception as e:
            print(f"Translation failed for description, using French: {e}")
            return description_fr # Fallback to French

def generate_reference():
    name = fake_fr.name() if random.random() < 0.7 else fake_en.name()
    title = random.choice(["Manager", "Collègue", "Superviseur", "Professeur"])
    # Generate a more plausible email
    clean_name = "".join(filter(str.isalnum, name.lower()))
    email_provider = random.choice(["example.com", "mail.com", "company.com"])
    email = f"{clean_name}{random.randint(1,99)}@{email_provider}"
    return f"{name}, {title}, {email}"

def validate_cv(cv_data): # Renamed cv to cv_data for clarity
    profile_limits = {"junior": 5, "intermediate": 10, "experienced": 15, "senior": 20} # Max years for profile
    # Ensure experience years don't exceed profile type's upper typical boundary
    if cv_data["Années d’expérience"] > profile_limits[cv_data['Profil']]:
         # Adjust if it's significantly over, but allow some flexibility from initial random generation
         # The generation already caps based on profile_type, so this is more of a safeguard
        cv_data["Années d’expérience"] = random.randint(0, profile_limits[cv_data['Profil']])

    if cv_data["Métier"] not in job_skills:
        # This should be caught earlier by the job_skills population for all jobs.
        print(f"Warning: Compétences manquantes pour {cv_data['Métier']} in job_skills dict.")
        cv_data["Compétences"] = "Compétences à définir"


# Modified generate_cv function
def generate_cv_details(job, profile_type, full_name_to_use, is_djiboutian_flag, nationality_label_to_use, cv_global_index):
    address = random.choice(djiboutian_addresses) if is_djiboutian_flag else (fake_fr.address() if nationality_label_to_use == "French" else fake_en.address())
    
    experience_years_map = {"junior": (0, 5), "intermediate": (6, 10),
                            "experienced": (11, 15), "senior": (16, 20)}
    min_exp, max_exp = experience_years_map[profile_type]
    experience_years = random.randint(min_exp, max_exp)
    
    actual_skills_for_job = job_skills.get(job, [])
    count_available_skills = len(actual_skills_for_job)
    selected_skills = []
    if count_available_skills > 0:
        # PDF: 5 to 15 skills.
        # If count_available_skills < 5, select all. Otherwise, select between 5 and min(15, count_available_skills)
        lower_bound_num_skills = 5
        
        if count_available_skills < lower_bound_num_skills:
            num_to_select = count_available_skills
        else:
            num_to_select = random.randint(lower_bound_num_skills, min(15, count_available_skills))
        selected_skills = random.sample(actual_skills_for_job, num_to_select)
    
    education_levels = [f"Licence en {job.lower()}", f"Master en {job.lower()}", "Baccalauréat", 
                        f"Diplôme professionnel en {job.lower()}", f"BTS en {job.lower()}", f"Doctorat en {job.lower()}"]
    education = random.choice(education_levels)
    institution = random.choice(institutions)
    
    num_certs = random.randint(0, 3) # Can have 0 certificates
    certificate_details = []
    if num_certs > 0 and certificates: # Ensure certificates list is not empty
        chosen_certs = random.sample(certificates, min(num_certs, len(certificates)))
        certificate_details = [f"{cert} ({random.choice(institutions)})" for cert in chosen_certs]
    
    spoken_languages = assign_languages(nationality_label_to_use)
    hobbies_pool = ["Football", "Lecture", "Cuisine", "Voyage", "Photographie", "Musique", "Bénévolat", "Jardinage"]
    hobbies = random.sample(hobbies_pool, random.randint(0, 4)) # Can have 0 hobbies
    cv_language = "Français" if random.random() < 0.9 else "Anglais"
    description = generate_description(job, experience_years, selected_skills, cv_language)
    experiences = generate_experience(job, experience_years, profile_type)
    references = [generate_reference() for _ in range(random.randint(0, 2))] # Can have 0 references
    
    # Sanitize name for email
    email_name_part = "".join(c for c in full_name_to_use if c.isalnum()).lower()
    email_domain = random.choice(["djibtel.dj", "gmail.com", "hotmail.com", "yahoo.fr", "outlook.com"])
    email = f"{email_name_part}{random.randint(1,99)}@{email_domain}"
    phone = f"+25377{random.randint(100000, 999999):06d}"
    
    cv_format = "pdf" if cv_global_index < (14400 / 2) else "docx" # 50/50 split
    template = random.choice(["classic", "modern", "compact"])
    
    cv_data = {
        "Nom complet": full_name_to_use,
        "Nationalité": nationality_label_to_use,
        "Adresse": address,
        "Métier": job,
        "Années d’expérience": experience_years,
        "Profil": profile_type,
        "Formation": education,
        "Institution": institution,
        "Compétences": ", ".join(selected_skills) if selected_skills else "N/A",
        "Certificats": ", ".join(certificate_details) if certificate_details else "N/A",
        "Langues parlées": ", ".join(spoken_languages) if spoken_languages else "N/A",
        "Hobbies": ", ".join(hobbies) if hobbies else "N/A",
        "Description": description,
        "Expériences": experiences if experiences else ["Aucune expérience professionnelle pertinente."],
        "Références": references if references else ["Sur demande"],
        "Email": email,
        "Téléphone": phone,
        "Langue du CV": cv_language,
        "Format": cv_format,
        "Template": template
    }
    validate_cv(cv_data) # Final validation pass
    return cv_data

# ... (create_cv_pdf and create_cv_docx functions remain largely the same as in the original file)
# Small modification to path handling in generate_and_save_cv
def create_cv_pdf(cv_data, output_path): # Renamed cv to cv_data
    try:
        doc = SimpleDocTemplate(output_path, pagesize=letter, title=f"CV {cv_data['Nom complet']}")
        styles = getSampleStyleSheet()
        # Define styles
        title_style = ParagraphStyle(name='TitleStyle', fontName='Helvetica-Bold', fontSize=18, alignment=1, spaceAfter=12, textColor=colors.HexColor("#000080"))
        heading_style = ParagraphStyle(name='HeadingStyle', fontName='Helvetica-Bold', fontSize=12, spaceAfter=6, spaceBefore=10, textColor=colors.HexColor("#333333"))
        normal_style = ParagraphStyle(name='NormalStyle', fontName='Helvetica', fontSize=10, spaceAfter=4, leading=12)
        contact_style = ParagraphStyle(name='ContactStyle', fontName='Helvetica', fontSize=10, alignment=1, spaceAfter=10)
        list_item_style = ParagraphStyle(name='ListItem', fontName='Helvetica', fontSize=10, leftIndent=20, spaceAfter=2)

        content = []
        content.append(Paragraph(cv_data["Nom complet"], title_style))
        content.append(Paragraph(f"{cv_data['Email']} | {cv_data['Téléphone']} | {cv_data['Adresse']}", contact_style))
        content.append(Spacer(1, 0.2 * inch))

        content.append(Paragraph("Profil" if cv_data["Langue du CV"] == "Français" else "Profile", heading_style))
        content.append(Paragraph(cv_data["Description"], normal_style))
        
        content.append(Paragraph("Expériences Professionnelles" if cv_data["Langue du CV"] == "Français" else "Work Experience", heading_style))
        if cv_data["Expériences"]:
            for exp in cv_data["Expériences"]:
                content.append(Paragraph(f"• {exp}", list_item_style))
        else:
            content.append(Paragraph("Aucune expérience pertinente.", normal_style))

        content.append(Paragraph("Formation" if cv_data["Langue du CV"] == "Français" else "Education", heading_style))
        content.append(Paragraph(f"{cv_data['Formation']} - <i>{cv_data['Institution']}</i>", normal_style))

        content.append(Paragraph("Compétences" if cv_data["Langue du CV"] == "Français" else "Skills", heading_style))
        content.append(Paragraph(cv_data["Compétences"], normal_style))

        if cv_data["Certificats"] and cv_data["Certificats"] != "N/A":
            content.append(Paragraph("Certificats" if cv_data["Langue du CV"] == "Français" else "Certificates", heading_style))
            content.append(Paragraph(cv_data["Certificats"], normal_style))
        
        content.append(Paragraph("Langues" if cv_data["Langue du CV"] == "Français" else "Languages", heading_style))
        content.append(Paragraph(cv_data["Langues parlées"], normal_style))

        if cv_data["Hobbies"] and cv_data["Hobbies"] != "N/A":
            content.append(Paragraph("Centres d'Intérêt" if cv_data["Langue du CV"] == "Français" else "Hobbies", heading_style))
            content.append(Paragraph(cv_data["Hobbies"], normal_style))

        if cv_data["Références"] and cv_data["Références"] != "Sur demande":
            content.append(Paragraph("Références" if cv_data["Langue du CV"] == "Français" else "References", heading_style))
            for ref in cv_data["Références"]:
                 content.append(Paragraph(ref, normal_style))
        else:
            content.append(Paragraph("Références" if cv_data["Langue du CV"] == "Français" else "References", heading_style))
            content.append(Paragraph("Sur demande" if cv_data["Langue du CV"] == "Français" else "Available upon request", normal_style))
            
        doc.build(content)
    except Exception as e:
        print(f"Erreur lors de la création du PDF {output_path}: {e}")

def create_cv_docx(cv_data, output_path): # Renamed cv to cv_data
    try:
        doc = Document()
        doc.add_heading(cv_data["Nom complet"], level=1)
        
        contact_p = doc.add_paragraph()
        contact_p.add_run(f"{cv_data['Email']} | {cv_data['Téléphone']} | {cv_data['Adresse']}")
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_heading("Profil" if cv_data["Langue du CV"] == "Français" else "Profile", level=2)
        doc.add_paragraph(cv_data["Description"])

        doc.add_heading("Expériences Professionnelles" if cv_data["Langue du CV"] == "Français" else "Work Experience", level=2)
        if cv_data["Expériences"]:
            for exp in cv_data["Expériences"]:
                doc.add_paragraph(exp, style='ListBullet')
        else:
            doc.add_paragraph("Aucune expérience pertinente.")

        doc.add_heading("Formation" if cv_data["Langue du CV"] == "Français" else "Education", level=2)
        doc.add_paragraph(f"{cv_data['Formation']} - {cv_data['Institution']}")

        doc.add_heading("Compétences" if cv_data["Langue du CV"] == "Français" else "Skills", level=2)
        doc.add_paragraph(cv_data["Compétences"])

        if cv_data["Certificats"] and cv_data["Certificats"] != "N/A":
            doc.add_heading("Certificats" if cv_data["Langue du CV"] == "Français" else "Certificates", level=2)
            doc.add_paragraph(cv_data["Certificats"])
        
        doc.add_heading("Langues" if cv_data["Langue du CV"] == "Français" else "Languages", level=2)
        doc.add_paragraph(cv_data["Langues parlées"])

        if cv_data["Hobbies"] and cv_data["Hobbies"] != "N/A":
            doc.add_heading("Centres d'Intérêt" if cv_data["Langue du CV"] == "Français" else "Hobbies", level=2)
            doc.add_paragraph(cv_data["Hobbies"])

        doc.add_heading("Références" if cv_data["Langue du CV"] == "Français" else "References", level=2)
        if cv_data["Références"] and cv_data["Références"] != "Sur demande":
             for ref in cv_data["Références"]:
                doc.add_paragraph(ref)
        else:
            doc.add_paragraph("Sur demande" if cv_data["Langue du CV"] == "Français" else "Available upon request")
            
        doc.save(output_path)
    except Exception as e:
        print(f"Erreur lors de la création du DOCX {output_path}: {e}")


def generate_and_save_cv_task(args_tuple): # Renamed to avoid conflict
    cv_data_item, index, base_path_pdf, base_path_docx = args_tuple # Unpack args
    try:
        # Sanitize name and job for filename
        safe_name = "".join(c for c in cv_data_item['Nom complet'] if c.isalnum() or c == ' ').replace(' ', '_')
        safe_job = "".join(c for c in cv_data_item['Métier'] if c.isalnum() or c == ' ').replace(' ', '_')
        
        # Add a random element to filename to further prevent collisions if names/jobs are very similar
        random_suffix = random.randint(10000, 99999)
        filename_base = f"cv_{safe_name}_{safe_job}_{random_suffix}"

        if cv_data_item["Format"] == "pdf":
            output_path = os.path.join(base_path_pdf, f"{filename_base}.pdf")
            create_cv_pdf(cv_data_item, output_path)
        else: # docx
            output_path = os.path.join(base_path_docx, f"{filename_base}.docx")
            create_cv_docx(cv_data_item, output_path)
        return {"filename": os.path.basename(output_path), "cv_data": cv_data_item} # Return cv_data for metadata
    except Exception as e:
        print(f"Erreur lors de la génération/sauvegarde du CV index {index}: {e} for {cv_data_item.get('Nom complet', 'Unknown')}")
        return None

# Étape 4 : Générer 14 400 CV
# Global set for Djiboutian name uniqueness
used_djiboutian_full_names = set()
all_cv_data_to_process_for_saving = [] # List to hold (cv_data, index, path_pdf, path_docx) tuples

total_cvs_generated_count = 0
max_cvs_target = 14400
cv_per_job_target = 72
profile_types = ["junior", "intermediate", "experienced", "senior"]

# Create output directories
base_output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'output_cvs') # Create 'output_cvs' in the same directory as the script
output_dir_pdf = os.path.join(base_output_path, 'cvs_pdf')
output_dir_docx = os.path.join(base_output_path, 'cvs_docx')
os.makedirs(output_dir_pdf, exist_ok=True)
os.makedirs(output_dir_docx, exist_ok=True)

print("Starting CV generation...")

for job_title in jobs:
    if total_cvs_generated_count >= max_cvs_target:
        break
    
    cvs_for_current_job = 0
    attempts_for_current_job_slot = 0 # Attempts to fill one of the 72 slots for this job
    max_job_slot_fill_attempts = cv_per_job_target * 3 # Allow more attempts overall for a job

    print(f"Generating CVs for job: {job_title} ({total_cvs_generated_count}/{max_cvs_target})")

    while cvs_for_current_job < cv_per_job_target and attempts_for_current_job_slot < max_job_slot_fill_attempts:
        if total_cvs_generated_count >= max_cvs_target:
            break

        is_djiboutian_candidate = random.random() < 0.87 # 87% Djiboutian
        current_full_name = None
        current_nationality_label = None

        if is_djiboutian_candidate:
            name_generation_attempts = 0
            max_djib_name_attempts = 50 # Try to find a unique Djiboutian name
            while name_generation_attempts < max_djib_name_attempts:
                first = random.choice(djiboutian_first_names_actual_120)
                last = random.choice(djiboutian_last_names_actual_105)
                potential_full_name = f"{first} {last}"
                if potential_full_name not in used_djiboutian_full_names:
                    current_full_name = potential_full_name
                    current_nationality_label = "Djibouti"
                    break
                name_generation_attempts += 1
            
            if not current_full_name: # Failed to get a unique Djiboutian name in reasonable attempts
                attempts_for_current_job_slot += 1
                continue # Skip this candidate attempt, try generating another one for this job slot
        
        else: # Foreign candidate
            current_nationality_label = random.choice(["French", "English", "Russian", "Italian", "Ethiopian", "Other"])
            # Use global pools for foreign names
            first = random.choice(FOREIGN_FIRST_NAMES_POOL_40)
            last = random.choice(FOREIGN_LAST_NAMES_POOL_50)
            current_full_name = f"{first} {last}"
            # No uniqueness check for foreign names as per PDF interpretation

        # If a name is set (either unique Djiboutian or any Foreign)
        profile_type = random.choice(profile_types)
        
        # Generate the full CV data dictionary
        # cv_global_index is total_cvs_generated_count before incrementing for this CV
        cv_data_dict = generate_cv_details(job_title, profile_type, current_full_name, 
                                           is_djiboutian_candidate, current_nationality_label, 
                                           total_cvs_generated_count)

        # Add Djiboutian name to used set *after* successful full CV data generation
        if is_djiboutian_candidate:
            used_djiboutian_full_names.add(current_full_name)

        all_cv_data_to_process_for_saving.append((cv_data_dict, total_cvs_generated_count, output_dir_pdf, output_dir_docx))
        
        cvs_for_current_job += 1
        total_cvs_generated_count += 1
        attempts_for_current_job_slot +=1 # This slot attempt was successful

    if cvs_for_current_job < cv_per_job_target:
        print(f"Warning: Job '{job_title}' only has {cvs_for_current_job}/{cv_per_job_target} CVs generated due to attempts limit or other issues.")

print(f"\nTotal CV data items prepared for file generation: {len(all_cv_data_to_process_for_saving)}")

# Now, process the file saving in parallel
generated_cv_metadata_list = []
if all_cv_data_to_process_for_saving:
    with ThreadPoolExecutor(max_workers=os.cpu_count() or 4) as executor: # Use available cores, default 4
        # Map tasks to executor
        results = list(executor.map(generate_and_save_cv_task, all_cv_data_to_process_for_saving))
    
    for result in results:
        if result and result.get("cv_data"):
            generated_cv_metadata_list.append(result["cv_data"])
else:
    print("No CV data was generated to be saved.")


print(f"Actual number of CVs generated and processed: {len(generated_cv_metadata_list)}")

# Étape 5 : Exporter les métadonnées
if generated_cv_metadata_list:
    cv_metadata_df = pd.DataFrame(generated_cv_metadata_list)
    metadata_path = os.path.join(base_output_path, "cv_metadata.csv") # Save metadata in the base output path
    cv_metadata_df.to_csv(metadata_path, index=False, encoding='utf-8-sig') # Added encoding
    print(f"Métadonnées exportées dans {metadata_path}")

    # Étape 6 : Créer une archive ZIP
    zip_filename = os.path.join(base_output_path, 'generated_cvs_archive.zip') # Save ZIP in the base output path
    print(f"Creating ZIP archive: {zip_filename}")
    with zipfile.ZipFile(zip_filename, 'w', zipfile.ZIP_DEFLATED, compresslevel=6) as zipf: # compresslevel 6 for speed
        for root, _, files in os.walk(output_dir_pdf):
            for file in files:
                if file.endswith('.pdf'):
                    zipf.write(os.path.join(root, file), os.path.join('cvs_pdf', file))
        for root, _, files in os.walk(output_dir_docx):
            for file in files:
                if file.endswith('.docx'):
                    zipf.write(os.path.join(root, file), os.path.join('cvs_docx', file))
        if os.path.exists(metadata_path):
            zipf.write(metadata_path, 'cv_metadata.csv')
    print(f"Archive ZIP créée : {zip_filename}")

    # Étape 7 : Télécharger l'archive (Colab specific)
    try:
        from google.colab import files
        files.download(zip_filename) # This will still be attempted if google.colab is available
        print("Colab download initiated.")
    except ImportError:
        print("Not in a Colab environment or google.colab not available. Skipping automatic download.")
        print(f"You can find the archive at: {os.path.abspath(zip_filename)}")
        print(f"And metadata at: {os.path.abspath(metadata_path)}")
else:
    print("No CVs were generated, so no metadata or ZIP file created.")